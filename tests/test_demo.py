# tests/test_parser_docx.py

import pytest
import docx
from app.utils.parser import extract_text_from_docx


def test_extract_text_from_docx_with_table(tmp_path):
    """
    测试：Word 深度解析是否能按顺序正确提取段落，并将表格转化为 Markdown
    """

    # ============================================================
    # 1. 动态造数据：利用 pytest 的 tmp_path 生成一个临时路径
    # ============================================================
    temp_file = tmp_path / "test_document.docx"

    # 使用 python-docx 动态创建一个包含段落和表格的 Word 文档
    doc = docx.Document()

    # 插入第一个段落
    doc.add_paragraph("第一章 交通违章处罚规定")

    # 插入一个 2行2列 的表格
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "违法行为"
    table.cell(0, 1).text = "罚款金额\n(人民币)"  # 故意加个换行符，测试我们的替换逻辑
    table.cell(1, 0).text = "机动车违规停放"
    table.cell(1, 1).text = "200元"

    # 插入第二个段落
    doc.add_paragraph("第二章 附则")

    # 保存这个临时 Word 文件
    doc.save(temp_file)

    # ============================================================
    # 2. 执行核心测试逻辑
    # ============================================================
    # 把临时文件的绝对路径传给我们的解析函数
    extracted_text = extract_text_from_docx(str(temp_file))

    # 打印出来方便在控制台查看（实际 pytest 运行时默认隐藏，除非失败或加 -s）
    print("\n=== 提取出的纯文本 ===")
    print(extracted_text)
    print("======================")

    # ============================================================
    # 3. 灵魂拷问 (断言 Assertions)
    # ============================================================
    # 拷问1：有没有正确注入页码？
    assert "<<PAGE:1>>" in extracted_text

    # 拷问2：段落是否被正确提取？
    assert "第一章 交通违章处罚规定" in extracted_text
    assert "第二章 附则" in extracted_text

    # 拷问3：【核心】表格是否被成功转换成了 Markdown 格式？
    # 并且换行符是否被替换成了空格（"罚款金额 (人民币)"）
    expected_table_header = "| 违法行为 | 罚款金额 (人民币) |"
    expected_table_divider = "|---|---|"
    expected_table_row = "| 机动车违规停放 | 200元 |"

    assert expected_table_header in extracted_text
    assert expected_table_divider in extracted_text
    assert expected_table_row in extracted_text

    # 拷问4：顺序是否正确？（段落1 -> 表格 -> 段落2）
    idx_p1 = extracted_text.find("第一章 交通违章处罚规定")
    idx_table = extracted_text.find("| 违法行为 |")
    idx_p2 = extracted_text.find("第二章 附则")

    assert idx_p1 < idx_table < idx_p2, "提取的顺序错乱了！"