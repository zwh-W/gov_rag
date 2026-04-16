# app/utils/parser.py
import re
import pdfplumber
import pytesseract
import docx
import os
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

from typing import List, Dict, Any

from app.core.config import settings
from app.core.logger import get_logger

logger = get_logger(__name__)

# ============================================================
# 【P0修改1】删除 MarkdownHeaderTextSplitter 的 import
#
# 原因：MarkdownHeaderTextSplitter 依赖文本里有 # ## ### 这种
#       Markdown 符号才能识别标题。
#       但政务 PDF 解析出来的文本长这样：
#           第一章 总则
#           第一条 为了规范...
#       完全没有 # 符号，导致整篇文档被当成一个 chunk，分块失效。
#
# 解决：改用正则表达式识别政务文档特有的"第X章/第X条/第X节"结构。
# ============================================================


# ============================================================
# 【P0修改2】新增政务文档结构识别的正则表达式
#
# 原因：政务文档有固定的层级结构：
#   第一层：第X章（最高级，如"第一章 总则"）
#   第二层：第X节（中间级，部分文档有，如"第一节 申请"）
#   第三层：第X条（最小可检索单元，如"第一条 为了规范..."）
#
# 为什么用中文数字+阿拉伯数字都匹配？
#   政务文档有的写"第一章"，有的写"第1章"，两种都要支持。
# ============================================================
_CHAPTER_PATTERN = re.compile(
    r'^第\s*[一二三四五六七八九十百千\d]+\s*章'
)
_SECTION_PATTERN = re.compile(
    r'^第\s*[一二三四五六七八九十百千\d]+\s*节'
)
_ARTICLE_PATTERN = re.compile(
    r'^第\s*[一二三四五六七八九十百千\d]+\s*条'
)


def extract_text_from_pdf(file_path: str) -> str:
    """
    从 PDF 中提取文本，包含扫描件 OCR 兜底。

    处理逻辑：
    - 正常 PDF：直接用 pdfplumber 提取文字
    - 扫描件（提取文字少于20字）：转图片后用 tesseract OCR
    """
    full_text = ""
    logger.info(f"开始解析PDF文件: {file_path}")
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text and len(page_text.strip()) > 20:
                    # ============================================================
                    # 【P0修改3】每页文本后面加页码标记
                    #
                    # 原因：后续分块时需要知道每个 chunk 来自第几页，
                    #       用于答案溯源（告诉用户"这段话来自第X页"）。
                    #       在文本里嵌入页码标记，分块函数可以提取它。
                    #
                    # 格式：<<PAGE:1>> 这种不会出现在正文里的特殊格式，
                    #       不会被正则误识别为章节标题。
                    # ============================================================
                    full_text += f"<<PAGE:{i + 1}>>\n{page_text}\n\n"
                else:
                    logger.info(f"第 {i + 1} 页文本量少，尝试 OCR...")
                    img = page.to_image(resolution=300).original
                    try:
                        ocr_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                        # OCR 结果同样加页码标记
                        full_text += f"<<PAGE:{i + 1}>>\n{ocr_text}\n\n"
                    except Exception as e:
                        logger.warning(f"第 {i + 1} 页 OCR 失败: {e}")
    except Exception as e:
        logger.error(f"解析 PDF 失败: {e}")
        raise
    return full_text


# ============================================================
# 【P2修改1】新增：Word 文档 (.docx) 提取逻辑
#
# 原因：系统允许上传 Word，但原来只有 PDF 解析器。
# 解决：使用 python-docx 提取文本。
# 注意：Word 是流式文本，没有严格的"物理页码"概念，
#       所以默认全部标记为第1页，重点保留段落和标题结构。
# ============================================================
def extract_text_from_docx(file_path: str) -> str:
    """从 Word 文档提取文本（段落 + 表格，保持原生顺序，Markdown表格）→ RAG 工业级"""
    full_text = "<<PAGE:1>>\n\n"
    logger.info(f"开始解析 Word 文件: {file_path}")

    try:
        doc = docx.Document(file_path)

        # ==============================
        # 核心：按文档原生顺序遍历（段落 + 表格）
        # ==============================
        for child in doc.element.body:
            # 1. 段落
            if isinstance(child, CT_P):
                para = Paragraph(child, doc)
                text = para.text.strip()
                if text:
                    full_text += text + "\n\n"

            # 2. 表格（输出标准 Markdown）
            elif isinstance(child, CT_Tbl):
                table = Table(child, doc)
                full_text += "\n"

                for row_idx, row in enumerate(table.rows):
                    row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                    full_text += "| " + " | ".join(row_data) + " |\n"

                    # 表头加分隔线
                    if row_idx == 0:
                        full_text += "|" + "|".join(["---"] * len(row.cells)) + "|\n"

                full_text += "\n"

        # ==============================
        # 增加：页眉页脚（政务必备）
        # ==============================
        full_text += "=== 页眉页脚信息 ===\n"
        for section in doc.sections:
            # 页眉
            for para in section.header.paragraphs:
                txt = para.text.strip()
                if txt:
                    full_text += f"页眉：{txt}\n"
            # 页脚
            for para in section.footer.paragraphs:
                txt = para.text.strip()
                if txt:
                    full_text += f"页脚：{txt}\n"

        full_text += "\n"

    except Exception as e:
        logger.error(f"解析 DOCX 失败: {str(e)}")
        raise

    return full_text


# ============================================================
# 【P2修改2】新增：统一的文档解析路由器
#
# 原因：让上游（processor）不需要关心底层是 PDF 还是 Word，
#       只管调用 extract_text，这里根据文件后缀自动路由。
#       这叫“策略模式 (Strategy Pattern)”的极简实现。
# ============================================================
def extract_text(file_path: str) -> str:
    """
    统一文档文本提取接口。
    根据文件后缀名自动调用对应的解析器。
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        # 注意：python-docx 原生只支持 .docx。
        # 如果是老版 .doc，在 Linux 环境下通常需要借助 libreoffice 转换。
        # 这里为了演示核心逻辑，统一视为 docx 处理。遇到老 doc 可能会抛异常。
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"不支持的文件格式解析: {ext}")





def chunk_text_by_headers(text: str) -> List[Dict[str, Any]]:
    """
    按政务文档的层级结构（章/节/条）进行智能分块。

    每个 chunk 包含：
    - content:    文本内容
    - breadcrumb: 面包屑路径，如 "第一章 总则 > 第三条"
    - page_number: 该 chunk 起始位置的页码
    """

    # ============================================================
    # 【P0修改4】完全重写分块逻辑，从 MarkdownHeaderTextSplitter
    #            改为基于正则的政务文档结构识别
    #
    # 核心思路：
    #   逐行扫描文本，遇到"第X章/节/条"就认为是一个新的分块边界。
    #   记录当前所在的章、节、条，拼接成面包屑路径存入每个 chunk。
    #
    # 为什么逐行而不是用正则直接 split？
    #   split 会把分隔符（章节标题）丢掉，我们需要把标题保留在
    #   chunk 内容里，让 chunk 本身是完整可读的。
    # ============================================================

    logger.info("开始政务文档结构化分块...")

    lines = text.split('\n')

    # 当前面包屑状态
    current_chapter = ""
    current_section = ""
    current_article = ""

    # 当前正在积累的 chunk
    current_content: List[str] = []
    current_page = 1
    chunk_start_page = 1  # 🚀【修复2】新增：专门记录当前块【开始】时的页码

    result_chunks: List[Dict[str, Any]] = []

    def _save_current_chunk():
        content = '\n'.join(current_content).strip()
        if not content:
            return

        breadcrumb_parts = [p for p in [current_chapter, current_section, current_article] if p]
        breadcrumb = " > ".join(breadcrumb_parts) if breadcrumb_parts else "正文"

        result_chunks.append({
            "content": content,
            "breadcrumb": breadcrumb,
            "page_number": chunk_start_page,  # 🚀【修复2】使用块开始时的页码，而不是乱变的当前页码
        })

    for line in lines:
        stripped = line.strip()

        if stripped.startswith('<<PAGE:') and stripped.endswith('>>'):
            try:
                current_page = int(stripped[7:-2])
            except ValueError:
                pass
            continue

            # 识别章级标题
        if _CHAPTER_PATTERN.match(stripped):
            _save_current_chunk()
            current_chapter = stripped
            current_section = ""
            current_article = ""
            # 🚀【修复1】章节标题只作为面包屑，不作为独立内容存入篮子！所以设为空列表！
            current_content = []

            # 识别节级标题
        elif _SECTION_PATTERN.match(stripped):
            _save_current_chunk()
            current_section = stripped
            current_article = ""
            # 🚀【修复1】同理，不作为独立内容
            current_content = []

        # 识别条级标题（最小切分单元）
        elif _ARTICLE_PATTERN.match(stripped):
            _save_current_chunk()
            current_article = stripped
            chunk_start_page = current_page  # 🚀【修复2】真正开始一个条文时，把当前的页码锁定给它！
            current_content = [line]

        # 普通正文行，追加到当前 chunk
        else:
            if not current_content:
                chunk_start_page = current_page  # 🚀【修复2】如果篮子是空的，装入第一句话时锁定页码
            current_content.append(line)

    # 循环结束后保存最后一个 chunk
    _save_current_chunk()

    # ============================================================
    # 【P0修改6】对过长的 chunk 进行二次切分
    #
    # 原因：有些"条"的内容非常长（比如附则里带大量列举项），
    #       一个 chunk 几千字会导致：
    #       1. 超过 Embedding 模型的最大输入长度（通常512 token）
    #       2. Rerank 模型输入超长被截断，精度下降
    #       3. 送给 LLM 的 context 太长，回答质量下降
    #
    # 解决：对超过 max_chunk_size 的 chunk 做滑动窗口二次切分，
    #       保留面包屑路径，加上序号区分。
    # ============================================================
    final_chunks = []
    max_chunk_size = settings.rag.chunk_size     # 字符数上限，可在 config.yaml 里配置
    overlap_size = settings.rag.chunk_overlap   # 重叠字符数，防止语义在边界断裂

    for chunk in result_chunks:
        content = chunk["content"]
        if len(content) <= max_chunk_size:
            final_chunks.append(chunk)
        else:
            # 滑动窗口切分
            sub_idx = 1
            start = 0
            while start < len(content):
                end = start + max_chunk_size
                sub_content = content[start:end]
                final_chunks.append({
                    "content": sub_content,
                    # 面包屑加上子块序号，方便溯源
                    "breadcrumb": f"{chunk['breadcrumb']} (片段{sub_idx})",
                    "page_number": chunk["page_number"],
                })
                sub_idx += 1
                start = end - overlap_size  # 下一个子块从重叠位置开始

    logger.info(f"分块完成，共生成 {len(final_chunks)} 个块（原始{len(result_chunks)}个，二次切分后{len(final_chunks)}个）")
    return final_chunks